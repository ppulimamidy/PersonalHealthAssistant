"""
Lab Results API — Phase 5: Lab Results & Health Twin

Manage lab test results, biomarker tracking, and AI-generated insights from lab data.
"""

# pylint: disable=too-many-locals,too-many-branches,too-many-statements,broad-except,line-too-long

import base64
import io
import json
import math
import os
import uuid
from datetime import date, datetime, timedelta, timezone
from typing import Any, Dict, List, Optional

import aiohttp
from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile
from pydantic import BaseModel, field_validator

from common.middleware.auth import get_current_user
from common.utils.logging import get_logger
from ..dependencies.usage_gate import (
    RateLimit,
    UsageGate,
    _supabase_get,
    _supabase_insert,
    _supabase_upsert,
    _supabase_patch,
    _supabase_delete,
)

logger = get_logger(__name__)
router = APIRouter()

# OpenAI for AI-generated insights
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")
OPENAI_MODEL = os.environ.get("LAB_INSIGHTS_AI_MODEL", "gpt-4o-mini")

# Anthropic for image scanning
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
ANTHROPIC_MODEL = "claude-sonnet-4-6"

# ---------------------------------------------------------------------------
# Pydantic Models
# ---------------------------------------------------------------------------


class Biomarker(BaseModel):
    """Individual biomarker measurement."""

    name: str
    value: float
    unit: str
    reference_range: Optional[str] = None
    status: str  # normal, borderline, abnormal, critical, unknown

    @field_validator("value")
    @classmethod
    def value_must_be_finite(cls, v: float) -> float:
        if not math.isfinite(v):
            raise ValueError(
                "Biomarker value must be a finite number (not NaN or infinity)"
            )
        return v


class CreateLabResultRequest(BaseModel):
    """Request to create a new lab result."""

    test_date: str  # YYYY-MM-DD
    test_type: str  # cbc, metabolic_panel, lipid_panel, etc.
    lab_name: Optional[str] = None
    ordering_provider: Optional[str] = None
    biomarkers: List[Biomarker]
    notes: Optional[str] = None
    tags: Optional[List[str]] = None


class LabResult(BaseModel):
    """Lab result with all biomarkers."""

    id: str
    user_id: str
    test_date: str
    test_type: str
    lab_name: Optional[str] = None
    ordering_provider: Optional[str] = None
    biomarkers: List[Biomarker]
    abnormal_count: int
    critical_count: int
    ai_summary: Optional[str] = None
    notes: Optional[str] = None
    tags: List[str]
    created_at: str
    updated_at: str
    # Anomaly detection fields — set on creation when abnormal/critical values are present
    anomaly_detected: bool = False
    anomaly_message: Optional[str] = None


class BiomarkerTrend(BaseModel):
    """Trend analysis for a single biomarker."""

    id: str
    biomarker_code: str
    biomarker_name: str
    trend_type: str  # improving, worsening, stable, fluctuating
    direction: str  # increasing, decreasing, stable
    slope: Optional[float] = None
    first_test_date: str
    last_test_date: str
    test_count: int
    current_value: float
    previous_value: Optional[float] = None
    min_value: float
    max_value: float
    avg_value: float
    std_deviation: float
    current_status: str
    trend_significance: str
    interpretation: str
    recommendations: List[str]
    computed_at: str


class LabInsight(BaseModel):
    """AI-generated insight from lab results."""

    id: str
    insight_type: str
    category: str
    title: str
    description: str
    key_findings: Dict[str, Any]
    correlated_health_metrics: Optional[Dict[str, float]] = None
    confidence_score: float
    priority: str
    recommendations: List[Dict[str, str]]
    is_acknowledged: bool
    created_at: str


class LabResultsResponse(BaseModel):
    """Response containing lab results."""

    results: List[LabResult]
    total_count: int


class BiomarkerTrendsResponse(BaseModel):
    """Response containing biomarker trends."""

    trends: List[BiomarkerTrend]
    generated_at: str


class LabInsightsResponse(BaseModel):
    """Response containing lab insights."""

    insights: List[LabInsight]
    total_count: int


class ScannedBiomarker(BaseModel):
    """A single biomarker extracted from a lab report image."""

    biomarker_code: str
    biomarker_name: str
    value: float
    unit: str
    reference_range: Optional[str] = None
    status: str = "normal"  # normal, borderline, abnormal, critical


class LabResultScanResult(BaseModel):
    """Result of scanning a lab report image."""

    test_type: Optional[str] = None
    test_date: Optional[str] = None
    lab_name: Optional[str] = None
    ordering_provider: Optional[str] = None
    biomarkers: List[ScannedBiomarker] = []
    notes: Optional[str] = None
    confidence: float = 0.0
    raw_text: Optional[str] = None


class LabProvider(BaseModel):
    """A lab provider that can be connected."""

    id: str
    name: str
    description: str
    logo: Optional[str] = None
    is_available: bool = False
    data_types: List[str] = []


# ---------------------------------------------------------------------------
# Helper Functions
# ---------------------------------------------------------------------------


def _compress_image(image_bytes: bytes, max_bytes: int = 4_500_000) -> bytes:
    """Compress image to fit within Anthropic's 5 MB limit."""
    if len(image_bytes) <= max_bytes:
        return image_bytes
    try:
        from PIL import Image  # type: ignore

        img = Image.open(io.BytesIO(image_bytes))
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        quality = 85
        scale = 1.0
        while True:
            buf = io.BytesIO()
            w = int(img.width * scale)
            h = int(img.height * scale)
            out = img.resize((w, h), Image.LANCZOS) if scale < 1.0 else img
            out.save(buf, format="JPEG", quality=quality, optimize=True)
            compressed = buf.getvalue()
            if len(compressed) <= max_bytes:
                return compressed
            if quality > 60:
                quality -= 10
            else:
                scale *= 0.75
            if scale < 0.1:
                return compressed
    except Exception:
        return image_bytes


# ---------------------------------------------------------------------------
# Biomarker reference ranges (module-level for efficiency)
# ---------------------------------------------------------------------------

# "High concern" ranges: (normal_min, normal_max, high_borderline, high_critical, unit)
# Abnormal: < normal_min  OR  > high_borderline
# Critical: > high_critical  OR  < normal_min (for biomarkers with a critical low not defined here)
# Borderline: normal_max < value <= high_borderline
_BIO_HIGH_BAD: Dict[str, tuple] = {
    # Lipids
    "total_cholesterol": (125, 200, 239, 400, "mg/dL"),
    "ldl": (0, 100, 159, 400, "mg/dL"),  # LDL has no clinical "critical" cutoff alone
    "triglycerides": (0, 150, 199, 500, "mg/dL"),
    "vldl": (2, 30, 40, 50, "mg/dL"),
    "non_hdl_cholesterol": (0, 130, 159, 190, "mg/dL"),
    # Blood sugar / diabetes
    "glucose": (65, 100, 125, 400, "mg/dL"),
    "hba1c": (4.0, 5.7, 6.4, 10.0, "%"),
    "insulin": (2, 25, 30, 100, "μIU/mL"),
    "c_peptide": (0.8, 3.1, 4.0, 7.0, "ng/mL"),
    # Kidney / uric acid
    "blood_urea_nitrogen": (7, 25, 30, 50, "mg/dL"),
    "creatinine": (0.6, 1.35, 1.7, 4.0, "mg/dL"),
    "uric_acid": (2.5, 7.0, 8.0, 12.0, "mg/dL"),
    # Liver enzymes
    "alanine_aminotransferase": (7, 56, 80, 200, "U/L"),
    "aspartate_aminotransferase": (10, 40, 60, 150, "U/L"),
    "alkaline_phosphatase": (44, 147, 200, 500, "U/L"),
    "gamma_glutamyltransferase": (5, 78, 100, 300, "U/L"),
    "total_bilirubin": (0.1, 1.2, 2.0, 5.0, "mg/dL"),
    "direct_bilirubin": (0.0, 0.3, 0.5, 1.5, "mg/dL"),
    "lactate_dehydrogenase": (140, 280, 350, 500, "U/L"),
    # Thyroid
    "tsh": (0.4, 4.5, 6.0, 20.0, "mIU/L"),
    # Inflammation
    "crp": (0, 3.0, 5.0, 10.0, "mg/L"),
    "hs_crp": (0, 1.0, 3.0, 10.0, "mg/L"),
    "esr": (0, 20, 40, 80, "mm/hr"),
    "homocysteine": (5, 15, 30, 50, "μmol/L"),
    "fibrinogen": (200, 400, 500, 700, "mg/dL"),
    # Cardiac
    "troponin_i": (0, 0.04, 0.1, 0.4, "ng/mL"),
    "bnp": (0, 100, 300, 900, "pg/mL"),
    "nt_pro_bnp": (0, 125, 400, 900, "pg/mL"),
    # WBC
    "white_blood_cells": (4.5, 11.0, 15.0, 30.0, "K/μL"),
    # Hormones
    "cortisol": (6, 23, 30, 60, "μg/dL"),
    # PSA
    "psa": (0, 4.0, 6.5, 10.0, "ng/mL"),
}

# Symmetric ranges: (critical_low, borderline_low, normal_min, normal_max, borderline_high, critical_high, unit)
# None = boundary not applicable in that direction
_BIO_SYMMETRIC: Dict[str, tuple] = {
    # CBC
    "hemoglobin": (7.0, 10.0, 11.5, 17.5, 18.5, 21.0, "g/dL"),
    "hematocrit": (20, 30, 35, 52, 55, 65, "%"),
    "red_blood_cells": (3.0, 3.5, 4.0, 6.2, 6.5, 7.0, "M/μL"),
    "platelets": (50, 100, 150, 400, 500, 1000, "K/μL"),
    "mcv": (60, 72, 80, 100, 110, 120, "fL"),
    "mch": (20, 24, 27, 33, 36, 42, "pg"),
    "mchc": (28, 30, 32, 36, 37, 40, "g/dL"),
    "neutrophils": (0.5, 1.5, 1.8, 7.7, 10, 20, "K/μL"),
    "lymphocytes": (0.5, 0.8, 1.0, 4.8, 6.0, 10.0, "K/μL"),
    "monocytes": (0, 0.1, 0.1, 1.0, 1.5, 3.0, "K/μL"),
    "eosinophils": (0, 0.0, 0.0, 0.5, 0.7, 1.5, "K/μL"),
    # Electrolytes  (Na≤130 severe hyponatremia; K>6.0 severe hyperkalemia)
    "sodium": (130, 131, 136, 145, 150, 160, "mEq/L"),
    "potassium": (2.5, 3.0, 3.5, 5.0, 5.5, 6.0, "mEq/L"),
    "chloride": (90, 95, 98, 107, 110, 120, "mEq/L"),
    "bicarbonate": (15, 18, 22, 29, 32, 40, "mEq/L"),
    "calcium": (6.5, 7.5, 8.5, 10.2, 10.8, 12.0, "mg/dL"),
    "total_protein": (4.0, 5.5, 6.3, 8.2, 8.8, 10.0, "g/dL"),
    "albumin": (2.0, 3.0, 3.5, 5.0, 5.3, 6.0, "g/dL"),
    "magnesium": (0.8, 1.4, 1.7, 2.2, 2.5, 3.0, "mg/dL"),
    "phosphorus": (1.5, 2.0, 2.5, 4.5, 5.0, 7.0, "mg/dL"),
    # Thyroid (both low and high matter)
    "free_t4": (0.4, 0.6, 0.8, 1.8, 2.0, 3.0, "ng/dL"),
    "free_t3": (1.5, 2.0, 2.3, 4.2, 4.8, 7.0, "pg/mL"),
    "total_t4": (2.0, 4.0, 5.0, 12.0, 13.0, 18.0, "μg/dL"),
    "total_t3": (60, 70, 80, 200, 220, 300, "ng/dL"),
    # GFR: <15 = kidney failure (critical), 15-50 = abnormal (CKD stage 3-4), 50-60 = borderline
    "gfr": (15, 50, 60, 150, None, None, "mL/min/1.73m²"),
    # Vitamins / minerals
    "vitamin_d": (0, 10, 30, 100, 125, 150, "ng/mL"),
    "vitamin_b12": (0, 150, 200, 900, 1200, 2000, "pg/mL"),
    "folate": (0, 2.0, 2.7, 17.0, 20.0, 30.0, "ng/mL"),
    "ferritin": (0, 7, 12, 336, 400, 1000, "ng/mL"),
    "serum_iron": (20, 45, 60, 170, 200, 350, "μg/dL"),
    "tibc": (150, 200, 250, 370, 400, 500, "μg/dL"),
    "transferrin_saturation": (0, 10, 20, 50, 55, 80, "%"),
    "zinc": (30, 50, 60, 120, 140, 200, "μg/dL"),
    # HDL: <20 = critical, 20-40 = abnormal (no borderline—below 40 is clinically significant)
    "hdl": (20, None, 40, 200, None, None, "mg/dL"),
    # Hormones
    "testosterone": (100, 200, 300, 1000, 1200, 2000, "ng/dL"),
    "dhea_s": (30, 65, 80, 560, 600, 800, "μg/dL"),
    "estradiol": (5, 10, 20, 200, 300, 500, "pg/mL"),
    "igf_1": (50, 100, 115, 307, 350, 500, "ng/mL"),
}

# Common aliases → canonical key used above
_BIO_ALIASES: Dict[str, str] = {
    # CBC
    "hgb": "hemoglobin",
    "hb": "hemoglobin",
    "hct": "hematocrit",
    "pcv": "hematocrit",
    "wbc": "white_blood_cells",
    "leukocytes": "white_blood_cells",
    "rbc": "red_blood_cells",
    "erythrocytes": "red_blood_cells",
    "plt": "platelets",
    "thrombocytes": "platelets",
    "neut": "neutrophils",
    "neutrophil": "neutrophils",
    "segs": "neutrophils",
    "lymphs": "lymphocytes",
    "lymph": "lymphocytes",
    "mono": "monocytes",
    "monos": "monocytes",
    "eos": "eosinophils",
    "eosino": "eosinophils",
    "baso": "basophils",
    # CMP
    "na": "sodium",
    "k": "potassium",
    "cl": "chloride",
    "co2": "bicarbonate",
    "hco3": "bicarbonate",
    "bicarb": "bicarbonate",
    "bun": "blood_urea_nitrogen",
    "urea": "blood_urea_nitrogen",
    "cr": "creatinine",
    "crea": "creatinine",
    "ca": "calcium",
    "tp": "total_protein",
    "protein_total": "total_protein",
    "alb": "albumin",
    "tbil": "total_bilirubin",
    "bilirubin": "total_bilirubin",
    "dbil": "direct_bilirubin",
    "alt": "alanine_aminotransferase",
    "sgpt": "alanine_aminotransferase",
    "ast": "aspartate_aminotransferase",
    "sgot": "aspartate_aminotransferase",
    "alp": "alkaline_phosphatase",
    "alk_phos": "alkaline_phosphatase",
    "ggt": "gamma_glutamyltransferase",
    "gamma_gt": "gamma_glutamyltransferase",
    "ldh": "lactate_dehydrogenase",
    "egfr": "gfr",
    "creatinine_clearance": "gfr",
    # Lipids
    "tc": "total_cholesterol",
    "chol": "total_cholesterol",
    "ldl_c": "ldl",
    "ldl_cholesterol": "ldl",
    "hdl_c": "hdl",
    "hdl_cholesterol": "hdl",
    "trig": "triglycerides",
    "trigs": "triglycerides",
    "tg": "triglycerides",
    # Thyroid
    "ft4": "free_t4",
    "t4_free": "free_t4",
    "ft3": "free_t3",
    "t3_free": "free_t3",
    "t4": "total_t4",
    "t3": "total_t3",
    # Vitamins
    "vit_d": "vitamin_d",
    "25_oh_vitamin_d": "vitamin_d",
    "25ohd": "vitamin_d",
    "vit_b12": "vitamin_b12",
    "b12": "vitamin_b12",
    "cobalamin": "vitamin_b12",
    "folic_acid": "folate",
    "ferr": "ferritin",
    "fe": "serum_iron",
    "iron": "serum_iron",
    "tsat": "transferrin_saturation",
    "mg": "magnesium",
    "zn": "zinc",
    "ua": "uric_acid",
    "urate": "uric_acid",
    # Inflammation
    "hscrp": "hs_crp",
    "high_sensitivity_crp": "hs_crp",
    "sed_rate": "esr",
    "homocys": "homocysteine",
    "fibrin": "fibrinogen",
    # Hormones
    "test": "testosterone",
    "total_testosterone": "testosterone",
    "e2": "estradiol",
    "dheas": "dhea_s",
    "cortisol_am": "cortisol",
    "igf1": "igf_1",
    "igf_i": "igf_1",
    # Cardiac
    "tni": "troponin_i",
    "troponin": "troponin_i",
    "ntprobnp": "nt_pro_bnp",
    "nt_bnp": "nt_pro_bnp",
    # Misc
    "a1c": "hba1c",
    "hemoglobin_a1c": "hba1c",
    "glycated_hemoglobin": "hba1c",
    "insulin_fasting": "insulin",
    "psa_total": "psa",
    "ph": "phosphorus",
}


def _analyze_biomarker_status(
    value: float,
    biomarker_code: str,
    unit: str,
    user_age: int = 35,
    user_gender: str = "all",
) -> tuple[str, str]:
    """
    Analyze biomarker status based on reference ranges.
    Returns (status, reference_range_str).
    Status: 'normal' | 'borderline' | 'abnormal' | 'critical' | 'unknown'
    """
    if not math.isfinite(value):
        return "unknown", "N/A"

    code = (
        biomarker_code.lower()
        .replace(" ", "_")
        .replace("-", "_")
        .replace("/", "_")
        .strip("_")
    )
    code = _BIO_ALIASES.get(code, code)

    # ── High-bad ranges ──────────────────────────────────────────────────────
    if code in _BIO_HIGH_BAD:
        norm_min, norm_max, border_hi, crit_hi, ref_unit = _BIO_HIGH_BAD[code]
        ref_range = f"{norm_min}–{norm_max} {ref_unit}"
        if norm_min is not None and value < norm_min:
            return "abnormal", ref_range
        if crit_hi is not None and value > crit_hi:
            return "critical", ref_range
        if value > norm_max:
            if border_hi is not None and value <= border_hi:
                return "borderline", ref_range
            return "abnormal", ref_range
        return "normal", ref_range

    # ── Symmetric ranges ─────────────────────────────────────────────────────
    if code in _BIO_SYMMETRIC:
        (
            crit_lo,
            border_lo,
            norm_min,
            norm_max,
            border_hi,
            crit_hi,
            ref_unit,
        ) = _BIO_SYMMETRIC[code]
        ref_range = f"{norm_min}–{norm_max} {ref_unit}"
        if crit_lo is not None and value <= crit_lo:
            return "critical", ref_range
        if norm_min is not None and value < norm_min:
            if border_lo is not None and value >= border_lo:
                return "borderline", ref_range
            return "abnormal", ref_range
        if crit_hi is not None and value >= crit_hi:
            return "critical", ref_range
        if norm_max is not None and value > norm_max:
            if border_hi is not None and value <= border_hi:
                return "borderline", ref_range
            return "abnormal", ref_range
        return "normal", ref_range

    # ── Unknown biomarker — do NOT silently default to "normal" ─────────────
    return "unknown", "Reference range not available"


async def _generate_lab_summary(
    biomarkers: List[Dict[str, Any]], test_type: str
) -> str:
    """Generate AI summary of lab results."""
    if not OPENAI_API_KEY:
        return _generate_simple_summary(biomarkers)

    abnormal_markers = [
        b for b in biomarkers if b.get("status") not in ("normal", "borderline")
    ]

    if not abnormal_markers:
        return "All biomarkers within normal ranges. Overall lab results look good."

    prompt = f"""You are a health data analyst. Summarize these lab test results for a patient.

Test Type: {test_type}
Abnormal Biomarkers:
{json.dumps(abnormal_markers, indent=2)}

Provide a brief 2-3 sentence summary highlighting:
1. Which biomarkers are abnormal and their significance
2. Any patterns or concerns
3. General recommendation (if any)

Keep it concise and patient-friendly. Do not provide medical diagnoses."""

    try:
        timeout = aiohttp.ClientTimeout(total=10)
        async with aiohttp.ClientSession(timeout=timeout) as session:
            async with session.post(
                "https://api.openai.com/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {OPENAI_API_KEY}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": OPENAI_MODEL,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.7,
                    "max_tokens": 150,
                },
            ) as resp:
                if resp.status != 200:
                    return _generate_simple_summary(biomarkers)

                result = await resp.json()
                summary = result["choices"][0]["message"]["content"].strip()
                return summary
    except Exception as exc:
        logger.warning("AI summary generation failed: %s", exc)
        return _generate_simple_summary(biomarkers)


def _generate_simple_summary(biomarkers: List[Dict[str, Any]]) -> str:
    """Generate simple text summary without AI."""
    abnormal = sum(1 for b in biomarkers if b.get("status") == "abnormal")
    critical = sum(1 for b in biomarkers if b.get("status") == "critical")

    if critical > 0:
        return f"{critical} critical value(s) detected. Please consult with your healthcare provider."
    elif abnormal > 0:
        return f"{abnormal} biomarker(s) outside normal range. Consider discussing with your doctor."
    else:
        return "All biomarkers within normal ranges."


def _calculate_biomarker_trend(
    values: List[float], dates: List[str]
) -> tuple[str, str, float]:
    """
    Calculate trend for biomarker values over time.
    Returns (trend_type, direction, slope).
    """
    if len(values) < 2:
        return "stable", "stable", 0.0

    # Simple linear regression
    n = len(values)
    x = [float(i) for i in range(n)]
    y = values

    mean_x = sum(x) / n
    mean_y = sum(y) / n

    numerator = sum((x[i] - mean_x) * (y[i] - mean_y) for i in range(n))
    denominator = sum((x[i] - mean_x) ** 2 for i in range(n))

    slope = numerator / denominator if denominator != 0 else 0.0

    # Determine direction
    if abs(slope) < 0.01:
        direction = "stable"
        trend_type = "stable"
    elif slope > 0:
        direction = "increasing"
        trend_type = "worsening" if slope > 0.1 else "fluctuating"
    else:
        direction = "decreasing"
        trend_type = "improving" if slope < -0.1 else "fluctuating"

    return trend_type, direction, slope


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------


@router.post("/lab-results", response_model=LabResult)
async def create_lab_result(
    request: CreateLabResultRequest,
    current_user: dict = Depends(UsageGate("lab_results")),
):
    """
    Create a new lab result entry.
    Requires Pro+ subscription.
    """
    user_id = current_user["id"]

    # Convert biomarkers to JSON format
    biomarkers_data = []
    abnormal_count = 0
    critical_count = 0

    for biomarker in request.biomarkers:
        # Analyze status if not provided
        status = biomarker.status
        ref_range = biomarker.reference_range

        if not ref_range:
            status, ref_range = _analyze_biomarker_status(
                biomarker.value, biomarker.name, biomarker.unit
            )

        biomarker_dict = {
            "name": biomarker.name,
            "value": biomarker.value,
            "unit": biomarker.unit,
            "reference_range": ref_range,
            "status": status,
        }
        biomarkers_data.append(biomarker_dict)

        if status == "abnormal":
            abnormal_count += 1
        elif status == "critical":
            critical_count += 1

    # Generate AI summary
    ai_summary = await _generate_lab_summary(biomarkers_data, request.test_type)

    # Create lab result
    lab_result = {
        "id": str(uuid.uuid4()),
        "user_id": user_id,
        "test_date": request.test_date,
        "test_type": request.test_type,
        "lab_name": request.lab_name,
        "ordering_provider": request.ordering_provider,
        "biomarkers": json.dumps(biomarkers_data),
        "abnormal_count": abnormal_count,
        "critical_count": critical_count,
        "ai_summary": ai_summary,
        "notes": request.notes,
        "tags": request.tags or [],
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }

    result = await _supabase_insert("lab_results", lab_result)

    if not result:
        raise HTTPException(status_code=500, detail="Failed to create lab result")

    # Parse biomarkers back for response
    result["biomarkers"] = json.loads(result["biomarkers"])

    # --- Anomaly detection ---
    anomaly_detected = False
    anomaly_message: Optional[str] = None

    if critical_count > 0:
        anomaly_detected = True
        anomaly_message = (
            f"{critical_count} critical value{'s' if critical_count > 1 else ''} detected. "
            "Consider logging how you're feeling — a symptoms and weight check-in can help correlate these results."
        )
    elif abnormal_count > 0:
        # Check if any of these biomarkers were also abnormal in a prior result
        prior_labs = await _supabase_get(
            "lab_results",
            f"user_id=eq.{user_id}&test_type=eq.{request.test_type}&order=test_date.desc&limit=3&select=biomarkers,test_date",
        )
        persistently_abnormal: list = []
        new_abnormal_names = {
            b["name"]
            for b in biomarkers_data
            if b.get("status") in ("abnormal", "critical")
        }
        for prior in prior_labs:
            prior_bm = prior.get("biomarkers")
            if isinstance(prior_bm, str):
                import json as _json

                try:
                    prior_bm = _json.loads(prior_bm)
                except Exception:
                    prior_bm = []
            if isinstance(prior_bm, list):
                for pb in prior_bm:
                    if pb.get("name") in new_abnormal_names and pb.get("status") in (
                        "abnormal",
                        "critical",
                    ):
                        persistently_abnormal.append(pb["name"])
                        break

        anomaly_detected = True
        if persistently_abnormal:
            anomaly_message = (
                f"{', '.join(set(persistently_abnormal))} {'remain' if len(set(persistently_abnormal)) > 1 else 'remains'} "
                "abnormal compared to prior results. Log a quick symptoms check-in to help track what's changed."
            )
        else:
            anomaly_message = (
                f"{abnormal_count} abnormal value{'s' if abnormal_count > 1 else ''} in these results. "
                "Log a symptoms check-in to capture how you're feeling."
            )

    result["anomaly_detected"] = anomaly_detected
    result["anomaly_message"] = anomaly_message

    return LabResult(**result)


@router.get("/lab-results", response_model=LabResultsResponse)
async def get_lab_results(
    test_type: Optional[str] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    limit: int = Query(default=50, le=100),
    current_user: dict = Depends(get_current_user),
):
    """Get user's lab results with optional filtering."""
    user_id = current_user["id"]

    # Build query
    query = f"user_id=eq.{user_id}"

    if test_type:
        query += f"&test_type=eq.{test_type}"
    if start_date:
        query += f"&test_date=gte.{start_date}"
    if end_date:
        query += f"&test_date=lte.{end_date}"

    query += f"&order=test_date.desc&limit={limit}"

    results = await _supabase_get("lab_results", query)

    if not results:
        return LabResultsResponse(results=[], total_count=0)

    # Parse biomarkers JSON
    for result in results:
        if isinstance(result.get("biomarkers"), str):
            result["biomarkers"] = json.loads(result["biomarkers"])

    return LabResultsResponse(
        results=[LabResult(**r) for r in results], total_count=len(results)
    )


@router.get("/lab-results/{result_id}", response_model=LabResult)
async def get_lab_result(
    result_id: str,
    current_user: dict = Depends(get_current_user),
):
    """Get a specific lab result by ID."""
    user_id = current_user["id"]

    results = await _supabase_get(
        "lab_results", f"id=eq.{result_id}&user_id=eq.{user_id}"
    )

    if not results:
        raise HTTPException(status_code=404, detail="Lab result not found")

    result = results[0]

    # Parse biomarkers JSON
    if isinstance(result.get("biomarkers"), str):
        result["biomarkers"] = json.loads(result["biomarkers"])

    return LabResult(**result)


@router.delete("/lab-results/{result_id}")
async def delete_lab_result(
    result_id: str,
    current_user: dict = Depends(get_current_user),
):
    """Delete a lab result."""
    user_id = current_user["id"]

    success = await _supabase_delete(
        "lab_results", f"id=eq.{result_id}&user_id=eq.{user_id}"
    )

    if not success:
        raise HTTPException(status_code=404, detail="Lab result not found")

    return {"success": True, "message": "Lab result deleted"}


@router.get("/biomarker-trends", response_model=BiomarkerTrendsResponse)
async def get_biomarker_trends(
    biomarker_code: Optional[str] = None,
    current_user: dict = Depends(UsageGate("lab_results")),
):
    """
    Get trend analysis for biomarkers.
    Analyzes all lab results to identify trends.
    Requires Pro+ subscription.
    """
    user_id = current_user["id"]

    # Fetch all lab results for the user
    results = await _supabase_get(
        "lab_results", f"user_id=eq.{user_id}&order=test_date.asc&limit=100"
    )

    if not results:
        return BiomarkerTrendsResponse(
            trends=[], generated_at=datetime.now(timezone.utc).isoformat()
        )

    # Parse biomarkers and group by biomarker name
    biomarker_data: Dict[str, List[Dict[str, Any]]] = {}

    for result in results:
        test_date = result["test_date"]
        biomarkers = (
            json.loads(result["biomarkers"])
            if isinstance(result["biomarkers"], str)
            else result["biomarkers"]
        )

        for bm in biomarkers:
            name = bm["name"]
            if biomarker_code and name.lower() != biomarker_code.lower():
                continue

            if name not in biomarker_data:
                biomarker_data[name] = []

            biomarker_data[name].append(
                {
                    "date": test_date,
                    "value": bm["value"],
                    "unit": bm["unit"],
                    "status": bm.get("status", "normal"),
                }
            )

    # Calculate trends
    trends = []

    for name, measurements in biomarker_data.items():
        if len(measurements) < 2:
            continue

        # Sort by date
        measurements.sort(key=lambda x: x["date"])

        values = [m["value"] for m in measurements]
        dates = [m["date"] for m in measurements]

        trend_type, direction, slope = _calculate_biomarker_trend(values, dates)

        current_value = values[-1]
        previous_value = values[-2] if len(values) >= 2 else None

        trend = {
            "id": str(uuid.uuid4()),
            "biomarker_code": name.lower().replace(" ", "_"),
            "biomarker_name": name,
            "trend_type": trend_type,
            "direction": direction,
            "slope": slope,
            "first_test_date": dates[0],
            "last_test_date": dates[-1],
            "test_count": len(values),
            "current_value": current_value,
            "previous_value": previous_value,
            "min_value": min(values),
            "max_value": max(values),
            "avg_value": sum(values) / len(values),
            "std_deviation": (
                sum((v - sum(values) / len(values)) ** 2 for v in values) / len(values)
            )
            ** 0.5,
            "current_status": measurements[-1]["status"],
            "trend_significance": "notable" if abs(slope) > 0.1 else "minor",
            "interpretation": f"{name} is {trend_type} ({direction})",
            "recommendations": [],
            "computed_at": datetime.now(timezone.utc).isoformat(),
        }

        trends.append(BiomarkerTrend(**trend))

    return BiomarkerTrendsResponse(
        trends=trends, generated_at=datetime.now(timezone.utc).isoformat()
    )


@router.get("/lab-insights", response_model=LabInsightsResponse)
async def get_lab_insights(
    priority: Optional[str] = None,
    current_user: dict = Depends(UsageGate("lab_results")),
):
    """
    Get AI-generated insights from lab results.
    Requires Pro+ subscription.
    """
    user_id = current_user["id"]

    query = f"user_id=eq.{user_id}"
    if priority:
        query += f"&priority=eq.{priority}"

    query += "&order=created_at.desc&limit=20"

    insights = await _supabase_get("lab_insights", query)

    if not insights:
        return LabInsightsResponse(insights=[], total_count=0)

    # Parse JSON fields
    for insight in insights:
        if isinstance(insight.get("key_findings"), str):
            insight["key_findings"] = json.loads(insight["key_findings"])
        if isinstance(insight.get("correlated_health_metrics"), str):
            insight["correlated_health_metrics"] = json.loads(
                insight["correlated_health_metrics"]
            )
        if isinstance(insight.get("recommendations"), str):
            insight["recommendations"] = json.loads(insight["recommendations"])

    return LabInsightsResponse(
        insights=[LabInsight(**i) for i in insights], total_count=len(insights)
    )


@router.patch("/lab-insights/{insight_id}/acknowledge")
async def acknowledge_insight(
    insight_id: str,
    current_user: dict = Depends(get_current_user),
):
    """Mark a lab insight as acknowledged."""
    user_id = current_user["id"]

    result = await _supabase_patch(
        "lab_insights",
        f"id=eq.{insight_id}&user_id=eq.{user_id}",
        {
            "is_acknowledged": True,
            "acknowledged_at": datetime.now(timezone.utc).isoformat(),
        },
    )

    if not result:
        raise HTTPException(status_code=404, detail="Insight not found")

    return {"success": True, "message": "Insight acknowledged"}


# ---------------------------------------------------------------------------
# Lab Report Image Scanning
# ---------------------------------------------------------------------------


_LAB_SCAN_PROMPT = """You are a medical laboratory report parser. Extract all information from this lab report into structured JSON.

Extract the following:
1. test_type: The name of the test panel (e.g. "Lipid Panel", "Complete Blood Count", "Comprehensive Metabolic Panel")
2. test_date: The date the test was performed (YYYY-MM-DD format, or null if not visible)
3. lab_name: The name of the laboratory or clinic (e.g. "Quest Diagnostics", "LabCorp")
4. ordering_provider: The ordering doctor's name (e.g. "Dr. Jane Smith") or null
5. biomarkers: An array of ALL biomarker results found, each with:
   - biomarker_code: Short code/abbreviation (e.g. "CHOL", "GLU", "HBA1C", "WBC"). Uppercase. Derive from name if not shown.
   - biomarker_name: Full name of the biomarker (e.g. "Total Cholesterol", "Glucose", "Hemoglobin A1c")
   - value: Numeric value as a float
   - unit: Unit of measurement (e.g. "mg/dL", "g/dL", "%", "mmol/L")
   - reference_range: Reference range string (e.g. "100-199 mg/dL", "<5.7%") or null
   - status: "normal", "borderline", "abnormal", or "critical" based on H/L/HH/LL flags or comparison to reference range
6. notes: Any relevant notes, flags, comments, or patient/accession info from the report
7. confidence: Your confidence score from 0.0 to 1.0 in the extraction accuracy

Return ONLY valid JSON matching this schema, no markdown, no explanation:
{
  "test_type": string|null,
  "test_date": string|null,
  "lab_name": string|null,
  "ordering_provider": string|null,
  "biomarkers": [...],
  "notes": string|null,
  "confidence": number
}"""


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract plain text from a .docx or .doc file."""
    try:
        import docx  # python-docx

        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table cell text (lab results are often in tables)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)
    except Exception as exc:
        logger.warning("DOCX text extraction failed: %s", exc)
        raise


def _build_claude_content(
    file_bytes: bytes, filename: str, content_type: str
) -> List[Dict[str, Any]]:
    """
    Build the Claude message content block(s) appropriate for the file type.

    - PDF  → native document block (Claude reads text and layout directly)
    - DOCX/DOC → extract text → send as plain-text prompt prefix
    - Image → image block (existing Vision path)
    """
    name_lower = (filename or "").lower()
    ct_lower = (content_type or "").lower()

    # ── PDF ──────────────────────────────────────────────────────────────────
    if "pdf" in ct_lower or name_lower.endswith(".pdf"):
        data_b64 = base64.standard_b64encode(file_bytes).decode()
        return [
            {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": data_b64,
                },
            },
            {"type": "text", "text": _LAB_SCAN_PROMPT},
        ]

    # ── Word document ─────────────────────────────────────────────────────────
    if (
        name_lower.endswith((".docx", ".doc"))
        or "word" in ct_lower
        or "officedocument" in ct_lower
    ):
        text = _extract_docx_text(file_bytes)
        combined = (
            "The following is the text content extracted from a Word document lab report.\n\n"
            f"{text}\n\n"
            f"{_LAB_SCAN_PROMPT}"
        )
        return [{"type": "text", "text": combined}]

    # ── Image (default) ───────────────────────────────────────────────────────
    compressed = _compress_image(file_bytes)
    data_b64 = base64.standard_b64encode(compressed).decode()
    image_mt = (
        ct_lower
        if ct_lower in ("image/jpeg", "image/png", "image/gif", "image/webp")
        else "image/jpeg"
    )
    return [
        {
            "type": "image",
            "source": {"type": "base64", "media_type": image_mt, "data": data_b64},
        },
        {"type": "text", "text": _LAB_SCAN_PROMPT},
    ]


def _repair_truncated_json(text: str) -> str:
    """
    Attempt to repair JSON that was cut off mid-stream (token limit hit).
    Strategy: find the last complete biomarker object `}` inside the
    biomarkers array, close the array, then close the outer object with
    sensible defaults for any missing top-level keys.
    """
    # Find the position of the last fully-closed biomarker object
    # by scanning backwards for a `}` that is followed only by
    # whitespace / commas / array/object closers.
    biomarkers_start = text.find('"biomarkers"')
    if biomarkers_start == -1:
        raise ValueError("No biomarkers key found")

    # Truncate at the last `}` that could end a biomarker object
    last_close = text.rfind("}")
    if last_close == -1:
        raise ValueError("No closing brace found")

    # Keep everything up to and including that last `}`
    truncated = text[: last_close + 1]

    # Close the biomarkers array and the outer object
    repaired = truncated.rstrip().rstrip(",") + "]}"

    # Verify it at least starts correctly
    if not repaired.lstrip().startswith("{"):
        raise ValueError("Cannot repair: does not start with {")

    return repaired


def _parse_scan_response(raw_text: str) -> LabResultScanResult:
    """Parse Claude's JSON response into a LabResultScanResult."""
    # Strip markdown code fences if present
    text = raw_text.strip()
    if text.startswith("```"):
        parts = text.split("```")
        text = parts[1] if len(parts) > 1 else text
        if text.startswith("json"):
            text = text[4:]
    text = text.strip()

    # Try clean parse first; fall back to truncation repair
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        logger.warning("JSON truncated — attempting repair")
        repaired = _repair_truncated_json(text)
        parsed = json.loads(repaired)

    biomarkers = []
    for b in parsed.get("biomarkers", []):
        try:
            biomarkers.append(
                ScannedBiomarker(
                    biomarker_code=str(b.get("biomarker_code", "")).upper() or "UNK",
                    biomarker_name=str(
                        b.get("biomarker_name", b.get("biomarker_code", "Unknown"))
                    ),
                    value=float(b.get("value", 0)),
                    unit=str(b.get("unit", "")),
                    reference_range=b.get("reference_range"),
                    status=str(b.get("status", "normal")).lower(),
                )
            )
        except Exception as be:
            logger.warning("Skipping malformed biomarker: %s — %s", b, be)

    return LabResultScanResult(
        test_type=parsed.get("test_type"),
        test_date=parsed.get("test_date"),
        lab_name=parsed.get("lab_name"),
        ordering_provider=parsed.get("ordering_provider"),
        biomarkers=biomarkers,
        notes=parsed.get("notes"),
        confidence=float(parsed.get("confidence", 0.7)),
        raw_text=text,
    )


@router.post("/scan-image", response_model=LabResultScanResult)
async def scan_lab_result_image(
    file: UploadFile = File(...),
    _rl: None = Depends(RateLimit("lab_scan", max_per_minute=5)),
    current_user: Dict[str, Any] = Depends(get_current_user),
) -> LabResultScanResult:
    """
    Extract lab results from an uploaded file.
    Supports: images (JPG/PNG/WEBP), PDF, and Word documents (DOCX/DOC).
    Uses Claude AI for extraction.
    """
    if not ANTHROPIC_API_KEY:
        raise HTTPException(
            status_code=503, detail="Lab report scanning not configured"
        )

    file_bytes = await file.read()
    filename = file.filename or ""
    content_type = file.content_type or ""

    logger.info(
        "Lab scan request: filename=%s content_type=%s size=%d",
        filename,
        content_type,
        len(file_bytes),
    )

    try:
        content_blocks = _build_claude_content(file_bytes, filename, content_type)
    except Exception as exc:
        raise HTTPException(
            status_code=422, detail=f"Could not read file: {exc}"
        ) from exc

    try:
        import certifi
        import ssl

        ssl_ctx = ssl.create_default_context(cafile=certifi.where())
        connector = aiohttp.TCPConnector(ssl=ssl_ctx)

        async with aiohttp.ClientSession(connector=connector) as session:
            payload: Dict[str, Any] = {
                "model": ANTHROPIC_MODEL,
                "max_tokens": 8192,
                "messages": [{"role": "user", "content": content_blocks}],
            }
            # PDF documents need the beta header
            headers: Dict[str, str] = {
                "x-api-key": ANTHROPIC_API_KEY,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json",
            }
            if any(b.get("type") == "document" for b in content_blocks):
                headers["anthropic-beta"] = "pdfs-2024-09-25"

            async with session.post(
                "https://api.anthropic.com/v1/messages",
                headers=headers,
                json=payload,
            ) as resp:
                if resp.status != 200:
                    body = await resp.text()
                    logger.error("Anthropic scan error %s: %s", resp.status, body)
                    raise HTTPException(status_code=502, detail="AI extraction failed")

                data = await resp.json()
                raw_text = data["content"][0]["text"]
                return _parse_scan_response(raw_text)

    except HTTPException:
        raise
    except json.JSONDecodeError as exc:
        logger.error("Failed to parse Claude response as JSON: %s", exc)
        raise HTTPException(status_code=422, detail="Could not parse lab report")
    except Exception as exc:
        logger.error("Lab scan failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"Scan failed: {exc}")


# ---------------------------------------------------------------------------
# Lab Provider Connections (stubs — third-party integrations)
# ---------------------------------------------------------------------------

_LAB_PROVIDERS = [
    LabProvider(
        id="labcorp",
        name="LabCorp",
        description="Access your LabCorp results directly. Covers hundreds of diagnostic tests.",
        is_available=False,
        data_types=["blood panels", "urinalysis", "genetic tests", "pathology"],
    ),
    LabProvider(
        id="quest",
        name="Quest Diagnostics",
        description="Connect your Quest Diagnostics account to auto-import lab results.",
        is_available=False,
        data_types=["comprehensive metabolic", "lipid panels", "CBC", "hormone panels"],
    ),
    LabProvider(
        id="health_gorilla",
        name="Health Gorilla",
        description="Aggregated lab data from hundreds of labs via Health Gorilla API.",
        is_available=False,
        data_types=["all major lab types", "multi-lab aggregation"],
    ),
    LabProvider(
        id="labcorp_ondemand",
        name="LabCorp On Demand",
        description="Order and receive direct-to-consumer lab tests without a doctor's order.",
        is_available=False,
        data_types=["direct-to-consumer panels", "wellness tests"],
    ),
]


@router.get("/providers", response_model=List[LabProvider])
async def get_lab_providers(
    current_user: Dict[str, Any] = Depends(get_current_user),
) -> List[LabProvider]:
    """Return the list of supported third-party lab providers."""
    return _LAB_PROVIDERS


@router.post("/connect-provider")
async def connect_lab_provider(
    provider_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user),
):
    """
    Initiate connection to a third-party lab provider.
    All integrations are not yet available — returns HTTP 501.
    """
    provider = next((p for p in _LAB_PROVIDERS if p.id == provider_id), None)
    if not provider:
        raise HTTPException(status_code=404, detail="Provider not found")

    raise HTTPException(
        status_code=501,
        detail={
            "coming_soon": True,
            "message": (
                f"{provider.name} direct integration is not yet available. "
                "In the meantime, use the 'Scan Lab Report' feature to upload your PDF or image."
            ),
        },
    )
